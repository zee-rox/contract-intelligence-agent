from pathlib import Path
from uuid import UUID

from app.ingestion.normalization import normalize_text
from app.schemas.documents import ExtractedParagraph
from app.schemas.sources import DocxSourceLocator


def parse_docx(path: Path, document_id: UUID) -> tuple[list[ExtractedParagraph], list[str]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX parsing") from exc

    try:
        document = Document(str(path))
    except Exception as exc:
        raise ValueError("invalid_docx") from exc

    paragraphs: list[ExtractedParagraph] = []
    warnings: list[str] = []
    paragraph_number = 1
    for paragraph in document.paragraphs:
        raw_text = paragraph.text
        normalized = normalize_text(raw_text)
        if not normalized:
            paragraph_number += 1
            continue
        paragraphs.append(
            ExtractedParagraph(
                paragraph_number=paragraph_number,
                section_number=1,
                style_name=paragraph.style.name if paragraph.style else None,
                raw_text=raw_text,
                normalized_text=normalized,
                source_locator=DocxSourceLocator(
                    section_number=1,
                    paragraph_start=paragraph_number,
                    paragraph_end=paragraph_number,
                    char_offset_start=0,
                    char_offset_end=len(normalized),
                ),
            )
        )
        paragraph_number += 1

    for table in document.tables:
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells if normalize_text(cell.text)]
            if not cells:
                continue
            text = " | ".join(cells)
            paragraphs.append(
                ExtractedParagraph(
                    paragraph_number=paragraph_number,
                    section_number=1,
                    style_name="Table",
                    raw_text=text,
                    normalized_text=text,
                    source_locator=DocxSourceLocator(
                        section_number=1,
                        paragraph_start=paragraph_number,
                        paragraph_end=paragraph_number,
                        char_offset_start=0,
                        char_offset_end=len(text),
                    ),
                )
            )
            paragraph_number += 1

    if not paragraphs:
        raise ValueError("no_extractable_content")
    return paragraphs, warnings
