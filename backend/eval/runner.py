from __future__ import annotations

import argparse
import hashlib
import json
import os
import tempfile
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document

from app.config import Settings
from app.llm.errors import LLMProviderError
from app.llm.factory import build_llm_provider
from app.services.analysis_service import AnalysisService
from app.services.document_service import DocumentService
from app.services.qa_service import QAService
from eval.metrics.common import f1, ratio
from eval.report import write_eval_markdown
from eval.schemas import EvaluationDataset, EvalDocument

DEFAULT_DATASET = Path(__file__).parent / "datasets" / "phase5_dataset.json"


def load_dataset(path: Path) -> EvaluationDataset:
    return EvaluationDataset.model_validate(json.loads(path.read_text(encoding="utf-8")))


def build_docx_payload(document: EvalDocument, directory: Path) -> bytes:
    path = directory / document.filename
    docx = Document()
    for paragraph in document.paragraphs:
        docx.add_heading(paragraph.heading, level=1)
        docx.add_paragraph(paragraph.text)
    docx.save(path)
    return path.read_bytes()


def run_evaluation(dataset_path: Path = DEFAULT_DATASET, output_dir: Path | None = None, eval_md_path: Path | None = None) -> dict[str, Any]:
    dataset = load_dataset(dataset_path)
    run_id = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    repo_root = Path(__file__).parents[2]
    output_dir = output_dir or Path(__file__).parent / "results"
    output_dir.mkdir(parents=True, exist_ok=True)
    result_path = output_dir / "latest.json"
    eval_md_path = eval_md_path or repo_root / "EVAL.md"
    start = time.perf_counter()

    with tempfile.TemporaryDirectory(prefix="contract-eval-") as temp:
        temp_path = Path(temp)
        settings = Settings(
            storage_root=temp_path / "storage",
            llm_provider="fake",
            retrieval_score_threshold=0.0,
        )
        document_service = DocumentService(settings)
        analysis_service = AnalysisService(settings)
        qa_service = QAService(settings)

        details: list[dict[str, Any]] = []
        for eval_doc in dataset.documents:
            document_start = time.perf_counter()
            payload = build_docx_payload(eval_doc, temp_path)
            ingestion = document_service.ingest_upload(
                eval_doc.filename,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                payload,
            )
            analysis = analysis_service.get_or_run_analysis(ingestion.document.document_id)
            qa_results = [
                {
                    "question": question.question,
                    "expected_refused": question.expected_refused,
                    "expected_answer_terms": question.expected_answer_terms,
                    "expected_citation_terms": question.expected_citation_terms,
                    "response": qa_service.answer_question(ingestion.document.document_id, question.question).model_dump(mode="json"),
                }
                for question in eval_doc.questions
            ]
            details.append(
                {
                    "dataset_document_id": eval_doc.document_id,
                    "runtime_seconds": round(time.perf_counter() - document_start, 4),
                    "source_sha256": hashlib.sha256(payload).hexdigest(),
                    "expected_clauses": [item.model_dump() for item in eval_doc.expected_clauses],
                    "chunks": [item.model_dump(mode="json") for item in ingestion.chunks],
                    "predicted_clauses": [item.model_dump(mode="json") for item in analysis.clauses],
                    "predicted_risks": [item.model_dump(mode="json") for item in analysis.risks],
                    "qa_results": qa_results,
                    "ocr_expected_pages": eval_doc.ocr_expected_pages,
                    "ocr_actual_pages": 0,
                }
            )

    metrics, error_analysis = compute_metrics(details)
    metrics["performance"] = {"total_seconds": round(time.perf_counter() - start, 4), "documents": len(details)}
    results = {
        "run_id": run_id,
        "dataset_version": dataset.dataset_version,
        "dataset_path": _display_path(dataset_path, repo_root),
        "result_path": _display_path(result_path, repo_root),
        "configuration": settings.safe_summary(),
        "metrics": metrics,
        "provider_comparison": provider_comparison(settings),
        "error_analysis": error_analysis,
        "documents": details,
    }
    result_path.write_text(json.dumps(results, indent=2, sort_keys=True), encoding="utf-8")
    write_eval_markdown(results, eval_md_path)
    return results


def compute_metrics(details: list[dict[str, Any]]) -> tuple[dict[str, Any], dict[str, Any]]:
    expected_clause_total = 0
    predicted_clause_total = 0
    clause_true_positive = 0
    boundary_supported = 0
    risk_total = 0
    risk_correct = 0
    citations_total = 0
    citations_valid = 0
    refusal_total = 0
    refusal_correct = 0
    weak_clause_categories: set[str] = set()
    boundary_failures: list[dict[str, str]] = []
    risk_disagreements: list[dict[str, str]] = []
    citation_failures: list[dict[str, str]] = []
    unsupported_answer_issues: list[dict[str, str]] = []
    incorrect_refusals: list[dict[str, str]] = []
    ocr_expected = 0
    ocr_actual = 0

    for detail in details:
        expected = detail["expected_clauses"]
        predicted = detail["predicted_clauses"]
        chunks = {chunk["chunk_id"]: chunk for chunk in detail["chunks"]}
        risks = {risk["clause_id"]: risk for risk in detail["predicted_risks"]}
        expected_clause_total += len(expected)
        predicted_clause_total += len(predicted)
        ocr_expected += int(detail["ocr_expected_pages"])
        ocr_actual += int(detail["ocr_actual_pages"])

        for expected_clause in expected:
            predicted_clause = next(
                (
                    item
                    for item in predicted
                    if item["clause_type"] == expected_clause["clause_type"]
                    and expected_clause["snippet"].lower() in item["clause_text"].lower()
                ),
                None,
            )
            if predicted_clause:
                clause_true_positive += 1
                boundary_supported += 1
                predicted_risk = risks.get(predicted_clause["clause_id"])
                if predicted_risk:
                    risk_total += 1
                    if predicted_risk["risk_level"] == expected_clause["risk_level"]:
                        risk_correct += 1
                    else:
                        risk_disagreements.append(
                            {
                                "document_id": detail["dataset_document_id"],
                                "clause_type": expected_clause["clause_type"],
                                "expected": expected_clause["risk_level"],
                                "actual": predicted_risk["risk_level"],
                            }
                        )
            else:
                weak_clause_categories.add(expected_clause["clause_type"])
                boundary_failures.append(
                    {
                        "document_id": detail["dataset_document_id"],
                        "clause_type": expected_clause["clause_type"],
                        "expected_snippet": expected_clause["snippet"],
                    }
                )

        for qa in detail["qa_results"]:
            response = qa["response"]
            refusal_total += 1
            if bool(response["refused"]) == bool(qa["expected_refused"]):
                refusal_correct += 1
            else:
                incorrect_refusals.append({"document_id": detail["dataset_document_id"], "question": qa["question"]})
            if qa["expected_refused"] and not response["refused"]:
                unsupported_answer_issues.append({"document_id": detail["dataset_document_id"], "question": qa["question"]})
            for citation in response["citations"]:
                citations_total += 1
                cited_chunk = chunks.get(citation["chunk_id"])
                valid = bool(
                    cited_chunk
                    and citation["quoted_snippet"]
                    and citation["quoted_snippet"] in cited_chunk["normalized_text"]
                )
                if valid:
                    citations_valid += 1
                else:
                    citation_failures.append({"document_id": detail["dataset_document_id"], "question": qa["question"]})

    precision = ratio(clause_true_positive, predicted_clause_total)
    recall = ratio(clause_true_positive, expected_clause_total)
    return (
        {
            "clause": {
                "precision": precision,
                "recall": recall,
                "f1": f1(precision, recall),
                "boundary_support_rate": ratio(boundary_supported, expected_clause_total),
            },
            "risk": {"accuracy": ratio(risk_correct, risk_total), "evaluated": risk_total},
            "citation": {"validity_rate": ratio(citations_valid, citations_total), "evaluated": citations_total},
            "refusal": {"accuracy": ratio(refusal_correct, refusal_total), "evaluated": refusal_total},
            "ocr": {
                "evaluated_pages": ocr_expected,
                "actual_ocr_pages": ocr_actual,
                "status": "not_evaluated_no_ocr_fixture" if ocr_expected == 0 else "evaluated",
            },
        },
        {
            "weak_clause_categories": sorted(weak_clause_categories),
            "boundary_failures": boundary_failures,
            "risk_disagreements": risk_disagreements,
            "unsupported_answer_issues": unsupported_answer_issues,
            "citation_failures": citation_failures,
            "incorrect_refusals": incorrect_refusals,
            "ocr_degradation": "not evaluated; dataset contains no OCR fixture" if ocr_expected == 0 else "see OCR metric",
        },
    )


def provider_comparison(settings: Settings) -> list[dict[str, str]]:
    comparison = [{"provider": "fake", "status": "completed", "notes": "Primary deterministic evaluation provider."}]
    if os.environ.get("EVAL_RUN_LLAMA_CPP") == "true":
        try:
            build_llm_provider(settings.model_copy(update={"llm_provider": "llamacpp"}))
            comparison.append({"provider": "llamacpp", "status": "configured", "notes": "Provider path constructed successfully."})
        except LLMProviderError as exc:
            comparison.append({"provider": "llamacpp", "status": "unavailable", "notes": str(exc)})
    else:
        comparison.append({"provider": "llamacpp", "status": "skipped", "notes": "Set EVAL_RUN_LLAMA_CPP=true and LLAMACPP_BASE_URL to run."})
    return comparison


def _display_path(path: Path, repo_root: Path) -> str:
    try:
        return str(path.resolve().relative_to(repo_root.resolve()))
    except ValueError:
        return str(path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Run the contract intelligence evaluation harness.")
    parser.add_argument("--dataset", type=Path, default=DEFAULT_DATASET)
    parser.add_argument("--output-dir", type=Path, default=Path(__file__).parent / "results")
    parser.add_argument("--eval-md", type=Path, default=Path(__file__).parents[2] / "EVAL.md")
    args = parser.parse_args()
    results = run_evaluation(args.dataset, args.output_dir, args.eval_md)
    print(json.dumps({"result_path": results["result_path"], "metrics": results["metrics"]}, indent=2, sort_keys=True))


if __name__ == "__main__":
    main()
