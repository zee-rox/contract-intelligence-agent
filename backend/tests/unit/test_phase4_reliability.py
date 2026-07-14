from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.config import Settings, get_settings
from app.llm.errors import LLMProviderError
from app.llm.gemini_provider import GeminiProvider
from app.schemas.analysis import AnalysisManifest, ClauseAnalysisResult
from app.schemas.clauses import ExtractedClause
from app.schemas.risks import RiskAssessment
from app.services import analysis_service, document_service
from app.services.analysis_service import AnalysisService
from app.services.document_service import DocumentService
from app.storage.atomic import atomic_write_text
from app.storage.repository import StorageRepository


def _docx_payload(tmp_path: Path) -> bytes:
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("Termination", level=1)
    document.add_paragraph("Either party may terminate with thirty days notice.")
    document.save(path)
    return path.read_bytes()


def test_interrupted_atomic_write_keeps_existing_file_and_removes_temp(tmp_path, monkeypatch) -> None:
    target = tmp_path / "artifact.json"
    target.write_text("old", encoding="utf-8")

    def fail_replace(source, destination) -> None:
        raise RuntimeError("interrupted")

    monkeypatch.setattr("app.storage.atomic.os.replace", fail_replace)

    with pytest.raises(RuntimeError):
        atomic_write_text(target, "new")

    assert target.read_text(encoding="utf-8") == "old"
    assert [path for path in tmp_path.iterdir() if path.name != "artifact.json"] == []


def test_gemini_provider_retries_timeouts_with_bound(monkeypatch) -> None:
    settings = Settings(llm_provider="gemini", google_api_key="secret", llm_max_retries=2, llm_timeout_seconds=0.1)
    provider = GeminiProvider(settings)
    calls = 0

    class TimeoutChat:
        def __init__(self, **kwargs: Any) -> None:
            self.kwargs = kwargs

        def invoke(self, messages: list[Any]) -> object:
            nonlocal calls
            calls += 1
            raise TimeoutError("timeout")

    def fake_chat_model(**kwargs: Any) -> TimeoutChat:
        assert kwargs["model"] == "gemini-2.5-flash"
        assert kwargs["api_key"] == "secret"
        assert kwargs["max_retries"] == 0
        return TimeoutChat(**kwargs)

    monkeypatch.setattr("app.llm.gemini_provider.ChatGoogleGenerativeAI", fake_chat_model)

    with pytest.raises(LLMProviderError):
        provider.generate([])

    assert calls == 3


def test_gemini_provider_returns_text_response(monkeypatch) -> None:
    class TextResponse:
        text = '{"clauses": []}'

    class TextChat:
        def __init__(self, **kwargs: Any) -> None:
            self.kwargs = kwargs

        def invoke(self, messages: list[Any]) -> object:
            return TextResponse()

    def fake_chat_model(**kwargs: Any) -> TextChat:
        return TextChat(**kwargs)

    monkeypatch.setattr("app.llm.gemini_provider.ChatGoogleGenerativeAI", fake_chat_model)
    provider = GeminiProvider(Settings(llm_provider="gemini", gemini_api_key="secret"))
    response = provider.generate([])

    assert response.provider == "gemini"
    assert response.content == '{"clauses": []}'


def test_gemini_provider_extracts_text_from_content_blocks(monkeypatch) -> None:
    class BlockResponse:
        text = ""
        content = [{"type": "text", "text": '{"clauses": []}'}, {"type": "text", "text": "\n"}]

    class BlockChat:
        def __init__(self, **kwargs: Any) -> None:
            self.kwargs = kwargs

        def invoke(self, messages: list[Any]) -> object:
            return BlockResponse()

    def fake_chat_model(**kwargs: Any) -> BlockChat:
        return BlockChat(**kwargs)

    monkeypatch.setattr("app.llm.gemini_provider.ChatGoogleGenerativeAI", fake_chat_model)
    provider = GeminiProvider(Settings(llm_provider="gemini", gemini_api_key="secret"))
    response = provider.generate([])

    assert response.content == '{"clauses": []}\n'


def test_safe_settings_summary_does_not_include_secret_values() -> None:
    settings = Settings(llm_api_key="llm-secret", google_api_key="google-secret", gemini_api_key="gemini-secret")
    summary = settings.safe_summary()

    assert "llm_api_key" not in summary
    assert "google_api_key" not in summary
    assert "gemini_api_key" not in summary
    assert "llm-secret" not in str(summary)
    assert "google-secret" not in str(summary)
    assert "gemini-secret" not in str(summary)


def test_failed_indexing_does_not_leave_ready_manifest(tmp_path, monkeypatch) -> None:
    def fail_index(*args, **kwargs) -> None:
        raise RuntimeError("index failed")

    monkeypatch.setattr(document_service.DocumentFaissIndex, "build_and_persist", fail_index)
    service = DocumentService(Settings(storage_root=tmp_path))

    with pytest.raises(RuntimeError):
        service.ingest_upload(
            "sample.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            _docx_payload(tmp_path),
        )

    manifests = list(tmp_path.glob("*/manifest.json"))
    assert len(manifests) == 1
    assert '"status": "failed"' in manifests[0].read_text(encoding="utf-8")


def test_corrupt_index_metadata_maps_to_typed_api_error(tmp_path, monkeypatch) -> None:
    get_settings.cache_clear()
    monkeypatch.setenv("STORAGE_ROOT", str(tmp_path))
    ingested = DocumentService(Settings(storage_root=tmp_path)).ingest_upload(
        "sample.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        _docx_payload(tmp_path),
    )
    repository = StorageRepository(tmp_path)
    repository.paths.index_metadata(ingested.document.document_id).write_text("{not-json", encoding="utf-8")
    client = TestClient(__import__("app.main", fromlist=["create_app"]).create_app())

    response = client.post(
        f"/documents/{ingested.document.document_id}/questions",
        json={"question": "How can a party terminate?"},
        headers={"x-request-id": "test-request"},
    )

    assert response.status_code == 500
    assert response.json()["error"]["code"] == "artifact_validation_failed"
    assert response.json()["error"]["request_id"] == "test-request"
    get_settings.cache_clear()


def test_duplicate_analysis_requests_run_graph_once(tmp_path, monkeypatch) -> None:
    ingested = DocumentService(Settings(storage_root=tmp_path)).ingest_upload(
        "sample.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        _docx_payload(tmp_path),
    )
    calls = 0
    document_id = ingested.document.document_id

    def fake_run(self, requested_document_id, chunks) -> ClauseAnalysisResult:
        nonlocal calls
        calls += 1
        clause = ExtractedClause(
            clause_id="clause_0000",
            document_id=requested_document_id,
            clause_type="termination",
            clause_heading="Termination",
            clause_text="Either party may terminate with thirty days notice.",
            source_chunk_ids=[chunks[0].chunk_id],
            source_locators=chunks[0].source_locators,
            confidence="high",
        )
        risk = RiskAssessment(
            clause_id="clause_0000",
            risk_level="low",
            risk_reason="No obvious baseline concern was detected.",
            observed_factors=["No obvious baseline concern was detected."],
            missing_expected_elements=[],
            confidence="medium",
            baseline_version="risk-baseline-v1",
        )
        manifest = AnalysisManifest(
            document_id=requested_document_id,
            status="completed",
            prompt_version="clause-extraction-v1",
            risk_baseline_version="risk-baseline-v1",
            provider="fake",
            model="fake",
        )
        return ClauseAnalysisResult(document_id=requested_document_id, clauses=[clause], risks=[risk], manifest=manifest)

    monkeypatch.setattr(analysis_service.AnalysisSupervisor, "run", fake_run)
    service = AnalysisService(Settings(storage_root=tmp_path))

    with ThreadPoolExecutor(max_workers=2) as executor:
        results = list(executor.map(lambda _: service.get_or_run_analysis(document_id), range(2)))

    assert calls == 1
    assert results[0].manifest.created_at == results[1].manifest.created_at
