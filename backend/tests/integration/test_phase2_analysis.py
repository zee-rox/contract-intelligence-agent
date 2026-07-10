from docx import Document

from app.config import Settings
from app.services.analysis_service import AnalysisService
from app.services.document_service import DocumentService


def test_analysis_persists_clauses_and_risks(tmp_path) -> None:
    docx_path = tmp_path / "analysis.docx"
    document = Document()
    document.add_heading("Liability", level=1)
    document.add_paragraph("Liability is unlimited for all damages.")
    document.save(docx_path)
    settings = Settings(storage_root=tmp_path)
    ingested = DocumentService(settings).ingest_upload(
        "analysis.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        docx_path.read_bytes(),
    )

    service = AnalysisService(settings)
    first = service.get_or_run_analysis(ingested.document.document_id)
    second = service.get_or_run_analysis(ingested.document.document_id)

    assert first.clauses
    assert first.risks
    assert first.risks[0].risk_level == "high"
    assert first.manifest.status == "completed"
    assert second.manifest.created_at == first.manifest.created_at
