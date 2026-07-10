import fitz
from docx import Document

from app.config import Settings
from app.services.document_service import DocumentService


def test_pdf_upload_persists_source_metadata(tmp_path) -> None:
    pdf_path = tmp_path / "sample.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Termination\nEither party may terminate on thirty days notice.")
    pdf.save(pdf_path)

    service = DocumentService(Settings(storage_root=tmp_path))
    response = service.ingest_upload("sample.pdf", "application/pdf", pdf_path.read_bytes())

    assert response.document.status == "ready"
    assert response.chunks[0].source_locators[0].source_type == "pdf"
    assert response.chunks[0].source_locators[0].page_number == 1


def test_docx_upload_uses_paragraph_locators(tmp_path) -> None:
    docx_path = tmp_path / "sample.docx"
    document = Document()
    document.add_heading("Confidentiality", level=1)
    document.add_paragraph("The parties shall keep confidential information confidential.")
    document.save(docx_path)

    service = DocumentService(Settings(storage_root=tmp_path))
    response = service.ingest_upload(
        "sample.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        docx_path.read_bytes(),
    )

    assert response.document.status == "ready"
    assert response.chunks[0].source_locators[0].source_type == "docx"
    assert response.chunks[0].source_locators[0].paragraph_start == 1
